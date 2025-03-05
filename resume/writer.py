"""
Handles writing formatted resumes
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_TAB_ALIGNMENT, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import logging
import shutil
import os
from resume.formatter import DocumentFormatter
from resume.parser import ResumeParser




class ResumeWriter:
    """Handles writing formatted resumes"""
    
    def __init__(self, formatter=None):
        self.formatter = formatter or DocumentFormatter()
        self.parser = ResumeParser()
    
    def document_writer(self, resume_text,output_path,resume_path):
        try:
            # Create a new document with the rewritten content
            new_doc = Document()
            
            # Set document margins for better space utilization
            sections = new_doc.sections
            for section in sections:
                section.top_margin = Inches(0.3)
                section.bottom_margin = Inches(0.3)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            
            # Add a style for paragraphs with reduced spacing
            styles = new_doc.styles
            style = styles.add_style('Compact', WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles['Normal']
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(3)
            style.paragraph_format.line_spacing = 1.0
            style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
            # Process the resume text
            lines = resume_text.strip().split('\n')
            
            # Extract name and contact info
            name = ""
            contact = ""
            for i, line in enumerate(lines):
                if line.strip() == "<NAME>":
                    name = lines[i+1].strip()
                if line.strip() == "<CONTACT>":
                    contact = lines[i+1].strip()
                    break
            
            # Add name and contact info to document
            name_para = new_doc.add_paragraph(style='Compact')
            name_run = name_para.add_run(name)
            name_run.bold = True
            name_run.font.size = Pt(16)
            name_run.font.name = "Calibri"

            # Add contact info
            contact_para = new_doc.add_paragraph(style='Compact')
            contact_run = contact_para.add_run(contact)
            contact_run.font.size = Pt(11)
            contact_run.font.name = "Calibri"
            
            # Add space after header
            spacing_para = new_doc.add_paragraph(style='Compact')
            spacing_para.paragraph_format.space_after = Pt(2)
            
            # Extract summary
            summary = ""
            for i, line in enumerate(lines):
                if line.strip() == "<SUMMARY>":
                    summary = lines[i+1].strip()
                    break
            
            # Add executive summary section
            self.formatter.add_section_heading(new_doc, "EXECUTIVE SUMMARY")
            summary_para = new_doc.add_paragraph(style='Compact')
            summary_run = summary_para.add_run(summary)
            summary_run.font.size = Pt(11)
            summary_run.font.name = "Calibri"
            
            # Extract skills
            skills = []
            in_skills_section = False
            for line in lines:
                if line.strip() == "<SKILLS>":
                    in_skills_section = True
                    continue
                if in_skills_section and line.strip() and not line.startswith("<"):
                    if "|" in line:
                        skills.extend([s.strip() for s in line.split("|")])
                    else:
                        skills.append(line.strip())
                if in_skills_section and line.startswith("<"):
                    in_skills_section = False
            
            # Add functional expertise section
            spacing_para = new_doc.add_paragraph(style='Compact')
            spacing_para.paragraph_format.space_after = Pt(2)
            self.formatter.add_section_heading(new_doc, "FUNCTIONAL EXPERTISE")
            
            # Calculate how many skills go in each column
            skills_per_column = (len(skills) + 1) // 2  # Ceiling division
            
            # Create skills in two columns using tab stops
            for i in range(skills_per_column):
                skills_para = new_doc.add_paragraph(style='Compact')
                
                # Set tab stop for second column
                skills_para.paragraph_format.tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.LEFT)
                
                # Add first column bullet and skill
                bullet_run = skills_para.add_run("• ")
                bullet_run.font.size = Pt(11)
                bullet_run.font.name = "Calibri"
                
                skill_run = skills_para.add_run(skills[i])
                skill_run.font.size = Pt(11)
                skill_run.font.name = "Calibri"
                
                # Add second column bullet and skill if available
                if i + skills_per_column < len(skills):
                    tab_run = skills_para.add_run("\t")
                    
                    bullet_run2 = skills_para.add_run("• ")
                    bullet_run2.font.size = Pt(11)
                    bullet_run2.font.name = "Calibri"
                    
                    skill_run2 = skills_para.add_run(skills[i + skills_per_column])
                    skill_run2.font.size = Pt(11)
                    skill_run2.font.name = "Calibri"
            
            # Process experience sections
            current_company = None
            current_position = None
            current_date = None
            company_positions = {}  # Dictionary to track positions at each company

            # First pass: collect all positions for each company
            for i, line in enumerate(lines):
                if line.strip() == "<COMPANY>":
                    current_company = lines[i+1].strip()
                    if current_company not in company_positions:
                        company_positions[current_company] = []
                
                elif line.strip() == "<POSITION_DATE>" and current_company:
                    position_date = lines[i+1].strip()
                    if "|" in position_date:
                        current_position, current_date = [p.strip() for p in position_date.split("|", 1)]
                        
                        # Collect bullets for this position
                        bullets = []
                        j = i + 2
                        while j < len(lines):
                            if lines[j].strip() == "<BULLET>":
                                bullets.append(lines[j+1].strip())
                                j += 2
                            elif lines[j].strip() in ["<COMPANY>", "<POSITION_DATE>", "<EDUCATION>", "<VOLUNTEERISM>"]:
                                break
                            else:
                                j += 1
                        
                        # Add position with its date and bullets
                        company_positions[current_company].append((current_position, current_date, bullets))

            # Second pass: output companies with all their positions
            for company, positions in company_positions.items():
                # Add company and location
                company_para = new_doc.add_paragraph(style='Compact')
                
                if "|" in company:
                    company_name, location = [c.strip() for c in company.split("|", 1)]
                    company_run = company_para.add_run(f"{company_name} | {location}")
                else:
                    company_run = company_para.add_run(company)
                
                company_run.bold = True
                company_run.font.size = Pt(12)
                company_run.font.name = "Calibri"
                
                # Set tab stop for date alignment
                tab_position = Inches(8.5 - 0.5)  # Page width minus right margin
                company_para.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT)
                
                # If multiple positions, add date range spanning all positions
                if len(positions) > 1:
                    logging.info(f"Processing multiple positions for company: {company}")
                    logging.info(f"Positions before sorting: {positions}")
                    
                    # Sort positions by date (assuming format MM/YYYY – MM/YYYY)
                    positions.sort(key=lambda x: self.parser.parse_date(x[1].split("–")[0].strip()), reverse=True)
                    logging.info(f"Positions after sorting: {positions}")
                    
                    # Fix: Parse dates more carefully to handle edge cases
                    all_start_dates = []
                    all_end_dates = []
                    all_start_tuples = []  # For accurate sorting
                    
                    logging.info("Extracting start and end dates from each position:")
                    for position, date_range, _ in positions:
                        logging.info(f"  Position: {position}, Date range: {date_range}")
                        if "–" in date_range:
                            parts = date_range.split("–")
                            if len(parts) == 2:
                                start_date = parts[0].strip()
                                end_date = parts[1].strip()
                                all_start_dates.append(start_date)
                                all_end_dates.append(end_date)
                                all_start_tuples.append((start_date, self.parser.parse_date(start_date)))
                                logging.info(f"    Extracted start date: {start_date}, end date: {end_date}")
                            else:
                                logging.warning(f"    Invalid date format: {date_range}")
                        else:
                            logging.warning(f"    No date separator found in: {date_range}")
                    
                    # Find earliest start date and latest end date
                    logging.info(f"All start dates: {all_start_dates}")
                    logging.info(f"All end dates: {all_end_dates}")
                    
                    if all_start_dates and all_end_dates:
                        # Sort dates using the parse_date function
                        all_start_tuples.sort(key=lambda x: x[1])  # Sort by the parsed tuple
                        earliest_date = all_start_tuples[0][0]  # Get the original date string
                        
                        # Sort end dates
                        end_date_tuples = [(date, self.parser.parse_date(date)) for date in all_end_dates]
                        end_date_tuples.sort(key=lambda x: x[1], reverse=True)
                        latest_date = end_date_tuples[0][0]  # Get the original date string
                        
                        date_range = f"{earliest_date} – {latest_date}"
                        logging.info(f"Final date range: {date_range}")
                        
                        date_run = company_para.add_run("\t" + date_range)
                        date_run.bold = True  # Make company-level date bold
                        date_run.font.size = Pt(11)
                        date_run.font.name = "Calibri"
                    else:
                        # Fallback to original method if parsing fails
                        logging.warning("Using fallback method for date extraction")
                        earliest_date = positions[-1][1].split("–")[0].strip()
                        latest_date = positions[0][1].split("–")[1].strip()
                        
                        logging.info(f"Fallback earliest date: {earliest_date}")
                        logging.info(f"Fallback latest date: {latest_date}")
                        
                        date_range = f"{earliest_date} – {latest_date}"
                        logging.info(f"Fallback date range: {date_range}")
                        
                        date_run = company_para.add_run("\t" + date_range)
                        date_run.bold = True  # Make company-level date bold
                        date_run.font.size = Pt(11)
                        date_run.font.name = "Calibri"
                else:
                    # Just one position, add its date
                    date_run = company_para.add_run("\t" + positions[0][1])
                    date_run.bold = True  # Make company-level date bold
                    date_run.font.size = Pt(11)
                    date_run.font.name = "Calibri"
                
                # Add each position
                for position, date, bullets in positions:
                    # Add position
                    position_para = new_doc.add_paragraph(style='Compact')
                    position_run = position_para.add_run(position)
                    position_run.italic = True
                    position_run.font.size = Pt(11)
                    position_run.font.name = "Calibri"
                    
                    # Add date for individual position if multiple positions
                    if len(positions) > 1:
                        position_para.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT)
                        date_run = position_para.add_run("\t" + date)
                        date_run.italic = True  # Keep italic
                        date_run.font.size = Pt(11)
                        date_run.font.name = "Calibri"
                    
                    # Add bullets for this position
                    for bullet_text in bullets:
                        self.formatter.add_bullet_point(new_doc, bullet_text)
            
            # Add education section
            spacing_para = new_doc.add_paragraph(style='Compact')
            spacing_para.paragraph_format.space_after = Pt(2)
            self.formatter.add_section_heading(new_doc, "EDUCATION")
            
            # Extract education info
            education_lines = []
            in_education = False
            for line in lines:
                if line.strip() == "<EDUCATION>":
                    in_education = True
                    continue
                if in_education and line.strip() and not line.startswith("<"):
                    education_lines.append(line.strip())
                if in_education and line.startswith("<") and line.strip() != "<EDUCATION>":
                    in_education = False
            
            # Create education table
            if education_lines:
                for i in range(0, len(education_lines), 3):
                    if i+2 < len(education_lines):
                        # Add institution and location with bullet
                        edu_para = new_doc.add_paragraph(style='Compact')
                        
                        # Add bullet
                        bullet_run = edu_para.add_run("• ")
                        bullet_run.font.size = Pt(11)
                        bullet_run.font.name = "Calibri"
                        
                        # Add institution and location
                        inst_run = edu_para.add_run(education_lines[i] + "  ")
                        inst_run.bold = True
                        inst_run.font.size = Pt(11)
                        inst_run.font.name = "Calibri"
                        
                        # Add date right-aligned
                        if "|" in education_lines[i+2]:
                            date = education_lines[i+2].split("|")[1].strip()
                            edu_para.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT)
                            date_run = edu_para.add_run("\t" + date)
                            date_run.bold = True  # Make education date bold
                            date_run.font.size = Pt(11)
                            date_run.font.name = "Calibri"
                        
                        # Add degree on next line with indentation
                        degree_para = new_doc.add_paragraph(style='Compact')
                        degree_para.paragraph_format.left_indent = Inches(0.25)
                        degree_run = degree_para.add_run(education_lines[i+1])
                        degree_run.font.size = Pt(11)
                        degree_run.font.name = "Calibri"
                        
                        # Add GPA on next line with indentation
                        if "GPA" in education_lines[i+2]:
                            gpa = education_lines[i+2].split("|")[0].strip() if "|" in education_lines[i+2] else education_lines[i+2]
                            gpa_para = new_doc.add_paragraph(style='Compact')
                            gpa_para.paragraph_format.left_indent = Inches(0.25)
                            gpa_run = gpa_para.add_run(gpa)
                            gpa_run.font.size = Pt(11)
                            gpa_run.font.name = "Calibri"
            
            # Add volunteerism section
            volunteerism_lines = []
            in_volunteerism = False
            for line in lines:
                if line.strip() == "<VOLUNTEERISM>":
                    in_volunteerism = True
                    continue
                if in_volunteerism and line.strip() and not line.startswith("<"):
                    volunteerism_lines.append(line.strip())
                if in_volunteerism and line.startswith("<") and line.strip() != "<VOLUNTEERISM>":
                    in_volunteerism = False
            
            if volunteerism_lines:
                spacing_para = new_doc.add_paragraph(style='Compact')
                spacing_para.paragraph_format.space_after = Pt(2)
                self.formatter.add_section_heading(new_doc, "VOLUNTEERISM")
                
                for i in range(0, len(volunteerism_lines), 2):
                    if i+1 < len(volunteerism_lines):
                        # Add organization and date with bullet
                        vol_para = new_doc.add_paragraph(style='Compact')
                        
                        # Add bullet
                        bullet_run = vol_para.add_run("• ")
                        bullet_run.font.size = Pt(11)
                        bullet_run.font.name = "Calibri"
                        
                        # Add organization and date
                        if "|" in volunteerism_lines[i]:
                            org, date = [v.strip() for v in volunteerism_lines[i].split("|", 1)]
                            org_run = vol_para.add_run(org + "  ")
                            org_run.bold = True
                            org_run.font.size = Pt(11)
                            org_run.font.name = "Calibri"
                            
                            vol_para.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT)
                            date_run = vol_para.add_run("\t" + date)
                            date_run.bold = True  # Make volunteerism date bold
                            date_run.font.size = Pt(11)
                            date_run.font.name = "Calibri"
                        else:
                            org_run = vol_para.add_run(volunteerism_lines[i])
                            org_run.bold = True
                            org_run.font.size = Pt(11)
                            org_run.font.name = "Calibri"
                        
                        # Add description on next line with indentation
                        desc_para = new_doc.add_paragraph(style='Compact')
                        desc_para.paragraph_format.left_indent = Inches(0.25)
                        desc_run = desc_para.add_run(volunteerism_lines[i+1])
                        desc_run.font.size = Pt(11)
                        desc_run.font.name = "Calibri"
            
            # Add other relevant information section
            other_info_lines = []
            in_other_info = False
            for line in lines:
                if line.strip() == "<OTHER RELEVANT INFORMATION>":
                    in_other_info = True
                    continue
                if in_other_info and line.strip() and not line.startswith("<"):
                    other_info_lines.append(line.strip())
                if in_other_info and line.startswith("<") and line.strip() != "<OTHER RELEVANT INFORMATION>":
                    in_other_info = False

            if other_info_lines:
                spacing_para = new_doc.add_paragraph(style='Compact')
                spacing_para.paragraph_format.space_after = Pt(2)
                self.formatter.add_section_heading(new_doc, "OTHER RELEVANT INFORMATION")
                
                # Process each line as a separate bullet point
                for line in other_info_lines:
                    if ":" in line:
                        # This line has a category and details
                        category, details = line.split(":", 1)
                        para = new_doc.add_paragraph(style='Compact')
                        para.paragraph_format.left_indent = Inches(0.25)
                        
                        # Add bullet
                        bullet_run = para.add_run("• ")
                        bullet_run.font.size = Pt(11)
                        bullet_run.font.name = "Calibri"
                        
                        # Add category in bold
                        cat_run = para.add_run(category.strip() + ": ")
                        cat_run.bold = True
                        cat_run.font.size = Pt(11)
                        cat_run.font.name = "Calibri"
                        
                        # Add details
                        details_run = para.add_run(details.strip())
                        details_run.font.size = Pt(11)
                        details_run.font.name = "Calibri"
                    else:
                        # Just a regular line
                        para = new_doc.add_paragraph(style='Compact')
                        para.paragraph_format.left_indent = Inches(0.25)
                        
                        # Add bullet
                        bullet_run = para.add_run("• ")
                        bullet_run.font.size = Pt(11)
                        bullet_run.font.name = "Calibri"
                        
                        # Add text
                        text_run = para.add_run(line.strip())
                        text_run.font.size = Pt(11)
                        text_run.font.name = "Calibri"
            
            # Save the document
            new_doc.save(output_path)
            return True
            
        except Exception as e:
            logging.error(f"Resume rewriting failed: {e}. Falling back to original resume.")
            # If there's an error, copy the original resume to the output path
            if os.path.exists(resume_path):
                shutil.copy(resume_path, output_path)
            return False



