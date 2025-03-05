"""
Handles document formatting operations
"""
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement



class DocumentFormatter:
    """Handles document formatting operations"""
    
    @staticmethod
    def add_section_heading(doc, text):
        """Add a section heading with consistent formatting and reduced spacing"""
        heading = doc.add_paragraph(style='Compact')
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading_run = heading.add_run(text)
        heading_run.bold = True
        heading_run.font.size = Pt(12)
        heading_run.font.name = "Calibri"
        heading_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
        # Add a horizontal line below the heading
        p = heading._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '999999')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        # Reduce space after heading
        heading.paragraph_format.space_after = Pt(3)
        
        return heading


    @staticmethod
    def add_bullet_point(doc, text, indent_level=0):
        """Add a bullet point with consistent formatting and reduced spacing"""
        bullet = doc.add_paragraph(style='Compact')  # Use compact style
        bullet.paragraph_format.left_indent = Inches(0.25 + (0.25 * indent_level))
        bullet.paragraph_format.first_line_indent = Inches(-0.25)
        
        # Check if the bullet point has a skill prefix (skill: text)
        if ':' in text:
            skill, achievement = text.split(':', 1)
            skill = skill.strip()
            achievement = achievement.strip()
            
            # Add bullet character
            bullet_char = bullet.add_run("• ")
            bullet_char.font.size = Pt(11)
            bullet_char.font.name = "Calibri"
            
            # Add skill in bold (not capitalized)
            skill_run = bullet.add_run(f"{skill.lower()}: ")
            skill_run.bold = True
            skill_run.font.size = Pt(11)
            skill_run.font.name = "Calibri"
            
            # Add achievement text
            achievement_run = bullet.add_run(achievement)
            achievement_run.font.size = Pt(11)
            achievement_run.font.name = "Calibri"
        else:
            # Add bullet character
            bullet_char = bullet.add_run("• ")
            bullet_char.font.size = Pt(11)
            bullet_char.font.name = "Calibri"
            
            # Just add the bullet text as is
            bullet_run = bullet.add_run(text)
            bullet_run.font.size = Pt(11)
            bullet_run.font.name = "Calibri"
        
        # Reduce space after bullet point
        bullet.paragraph_format.space_after = Pt(2)
        
        return bullet


