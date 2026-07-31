"""Extracting text from a DOCX the way an ATS does, not the way Word displays it.

The gap between those two is where resumes die. Word renders headers, text
boxes, and multi-column tables into a layout a human reads correctly. A resume
parser walks the document body in XML order and takes what it finds, which
means content can be silently dropped or scrambled before any human sees it.

This module produces both readings:

* :func:`human_text` is everything with text in it, in document order,
  including the parts Word draws outside the main body flow.
* :func:`ats_text` simulates a mainstream parser: body paragraphs and tables
  only, tables flattened cell by cell in row-major order, with headers,
  footers, and text boxes dropped.

Whatever appears in the first and not the second is content the candidate
believes they submitted and the employer never receives.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import List

from docx import Document
from docx.oxml.ns import qn

# Word stores drawing-canvas text in these elements. They render on screen and
# are routinely skipped by parsers that only walk w:body paragraphs. The
# markup-compatibility namespace is not in python-docx's prefix map, so it is
# spelled out rather than resolved through qn().
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
TEXTBOX_TAGS = (qn("w:txbxContent"), f"{{{MC_NS}}}AlternateContent")


@dataclass
class ExtractionReport:
    """Both readings of one document, plus what differs between them."""

    ats_text: str
    human_text: str
    body_paragraphs: int = 0
    table_count: int = 0
    table_cell_texts: List[str] = field(default_factory=list)
    header_texts: List[str] = field(default_factory=list)
    footer_texts: List[str] = field(default_factory=list)
    textbox_texts: List[str] = field(default_factory=list)
    image_count: int = 0
    hyperlink_targets: List[str] = field(default_factory=list)

    @property
    def dropped_text(self) -> List[str]:
        """Text a human sees that the simulated parser never receives."""
        return [
            snippet
            for snippet in self.header_texts + self.footer_texts + self.textbox_texts
            if snippet.strip()
        ]

    @property
    def has_tables(self) -> bool:
        return self.table_count > 0


def _iter_block_items(parent):
    """Yield paragraphs and tables from the body in true document order.

    ``Document.paragraphs`` and ``Document.tables`` are separate sequences, so
    neither preserves how the two interleave. Reading order is exactly what a
    parser gets wrong, so the order has to come from the XML.
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = parent.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _textbox_strings(document) -> List[str]:
    """Every run of text living inside a text box or drawing canvas."""
    found: List[str] = []
    for tag in TEXTBOX_TAGS:
        for node in document.element.body.iter(tag):
            text = "".join(t.text or "" for t in node.iter(qn("w:t")))
            if text.strip():
                found.append(text.strip())
    return found


def _section_strings(document, part: str) -> List[str]:
    """Text from every section's header or footer."""
    found: List[str] = []
    for section in document.sections:
        container = getattr(section, part, None)
        if container is None:
            continue
        for paragraph in container.paragraphs:
            if paragraph.text.strip():
                found.append(paragraph.text.strip())
        for table in getattr(container, "tables", []):
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        found.append(cell.text.strip())
    return found


def extract(path: str | Path) -> ExtractionReport:
    """Read a DOCX twice: once as a parser sees it, once as a human does."""
    document = Document(str(path))

    ats_parts: List[str] = []
    table_cells: List[str] = []
    body_paragraphs = 0
    table_count = 0

    for block in _iter_block_items(document):
        if hasattr(block, "rows"):  # Table
            table_count += 1
            # Row-major flattening is what turns a two-column layout into
            # interleaved nonsense: "Skills Python 2019 Company".
            for row in block.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        table_cells.append(text)
                        ats_parts.append(text)
        else:
            text = block.text.strip()
            if text:
                body_paragraphs += 1
                ats_parts.append(text)

    headers = _section_strings(document, "header")
    footers = _section_strings(document, "footer")
    textboxes = _textbox_strings(document)

    images = len(document.element.body.findall(f".//{qn('a:blip')}"))
    links = [
        rel.target_ref
        for rel in document.part.rels.values()
        if "hyperlink" in rel.reltype
    ]

    ats_text = "\n".join(ats_parts)
    human_text = "\n".join(headers + ats_parts + textboxes + footers)

    return ExtractionReport(
        ats_text=ats_text,
        human_text=human_text,
        body_paragraphs=body_paragraphs,
        table_count=table_count,
        table_cell_texts=table_cells,
        header_texts=headers,
        footer_texts=footers,
        textbox_texts=textboxes,
        image_count=images,
        hyperlink_targets=links,
    )
