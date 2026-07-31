"""Synthetic resumes with known, deliberately planted defects.

A detector nobody has tested against a document whose problems are known in
advance is a detector nobody has tested. Each builder here produces a DOCX with
a documented set of flaws, so the checks can be verified as finding what is
there and, just as importantly, not inventing what is not.

All fixtures are generated from code, so nothing personal or copyrighted is
committed to the repository and the corpus rebuilds identically anywhere.
"""

from __future__ import annotations

from pathlib import Path
from typing import Callable, Dict, List

from docx import Document
from docx.shared import Inches, Pt

CONTACT = "jordan.reyes@example.com | (555) 123-4567 | Austin, TX"

EXPERIENCE = [
    (
        "Senior Data Engineer, Northwind Analytics",
        "Jan 2021 - Present",
        [
            "Built streaming pipelines in Python and Apache Spark processing 4TB daily",
            "Migrated batch ETL to Airflow, cutting nightly runtime from 6 hours to 90 minutes",
            "Designed dimensional models in Snowflake supporting 40 analysts",
            "Introduced unit testing and CI/CD via GitHub Actions across 12 repositories",
        ],
    ),
    (
        "Data Engineer, Cobalt Systems",
        "Jun 2018 - Dec 2020",
        [
            "Developed REST APIs in FastAPI serving machine learning predictions",
            "Automated data quality checks in SQL, catching 200+ defects before release",
            "Partnered with product management on cross functional analytics projects",
        ],
    ),
]

EDUCATION = "B.S. Computer Science, University of Texas at Austin, 2018"

SKILLS = (
    "Python, SQL, Apache Spark, Airflow, Snowflake, AWS, Docker, Kubernetes, "
    "FastAPI, Git, CI/CD, machine learning, data modeling, ETL"
)

SUMMARY = (
    "Data engineer with 6 years building distributed systems and analytics "
    "infrastructure. Focused on reliability, testing, and making data trustworthy "
    "enough that people act on it."
)

JOB_POSTING = """Senior Data Engineer

We are seeking a Senior Data Engineer to join our platform team.

Requirements:
- 5+ years of experience with Python and SQL
- Strong experience with Apache Spark and distributed systems
- Experience with Airflow or similar orchestration tools
- Proficiency in cloud platforms, particularly AWS
- Experience with Snowflake or comparable data warehouses
- Familiarity with Docker and Kubernetes
- Knowledge of CI/CD practices and unit testing

Preferred:
- Experience with machine learning pipelines
- Background in dimensional modeling and ETL design
- Exposure to Kafka and real-time streaming
- Terraform and infrastructure as code

You will build and maintain data pipelines, partner with cross functional teams,
and improve the reliability of our data platform.
"""


def _add_body(doc: Document, heading_style: Dict[str, str] | None = None) -> None:
    """Write a clean, conventional resume body into ``doc``."""
    labels = heading_style or {
        "summary": "PROFESSIONAL SUMMARY",
        "experience": "PROFESSIONAL EXPERIENCE",
        "education": "EDUCATION",
        "skills": "TECHNICAL SKILLS",
    }

    doc.add_paragraph(labels["summary"]).runs[0].bold = True
    doc.add_paragraph(SUMMARY)

    doc.add_paragraph(labels["experience"]).runs[0].bold = True
    for title, dates, bullets in EXPERIENCE:
        line = doc.add_paragraph()
        line.add_run(title).bold = True
        line.add_run(f"    {dates}")
        for bullet in bullets:
            doc.add_paragraph(bullet, style="List Bullet")

    doc.add_paragraph(labels["education"]).runs[0].bold = True
    doc.add_paragraph(EDUCATION)

    doc.add_paragraph(labels["skills"]).runs[0].bold = True
    doc.add_paragraph(SKILLS)


def build_clean(path: Path) -> List[str]:
    """A well-formed resume. Should score high and trip nothing critical."""
    doc = Document()
    doc.add_paragraph("JORDAN REYES").runs[0].bold = True
    doc.add_paragraph(CONTACT)
    _add_body(doc)
    doc.save(str(path))
    return []


def build_header_contact(path: Path) -> List[str]:
    """Contact details in the page header, where parsers do not look."""
    doc = Document()
    header = doc.sections[0].header
    header.paragraphs[0].text = f"JORDAN REYES | {CONTACT}"
    doc.add_paragraph("PROFILE").runs[0].bold = True
    _add_body(doc)
    doc.save(str(path))
    return ["dropped_content", "contact_details"]


def build_table_layout(path: Path) -> List[str]:
    """A two-column table layout, which flattens into interleaved nonsense.

    Also carries literal bullet glyphs, since hand-typed "•" is what people
    reach for once a table has taken their list styles away.
    """
    doc = Document()
    doc.add_paragraph("JORDAN REYES").runs[0].bold = True
    doc.add_paragraph(CONTACT)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    left, right = table.rows[0].cells
    left.width = Inches(2.0)
    right.width = Inches(4.5)

    left.paragraphs[0].text = "TECHNICAL SKILLS"
    left.add_paragraph(SKILLS)
    left.add_paragraph("EDUCATION")
    left.add_paragraph(EDUCATION)

    right.paragraphs[0].text = "PROFESSIONAL EXPERIENCE"
    for title, dates, bullets in EXPERIENCE:
        right.add_paragraph(f"{title}    {dates}")
        for bullet in bullets:
            right.add_paragraph(f"• {bullet}")

    doc.save(str(path))
    return ["table_layout", "risky_characters"]


def build_unlabeled_sections(path: Path) -> List[str]:
    """Creative section names that map to no known field."""
    doc = Document()
    doc.add_paragraph("JORDAN REYES").runs[0].bold = True
    doc.add_paragraph(CONTACT)
    _add_body(
        doc,
        {
            "summary": "WHO I AM",
            "experience": "WHERE I'VE BEEN",
            "education": "HOW I LEARNED IT",
            "skills": "WHAT I'M GOOD AT",
        },
    )
    doc.save(str(path))
    return ["section_headings"]


def build_no_dates(path: Path) -> List[str]:
    """Vague date language a parser cannot convert to a duration."""
    doc = Document()
    doc.add_paragraph("JORDAN REYES").runs[0].bold = True
    doc.add_paragraph(CONTACT)
    doc.add_paragraph("PROFESSIONAL SUMMARY").runs[0].bold = True
    doc.add_paragraph(SUMMARY)
    doc.add_paragraph("PROFESSIONAL EXPERIENCE").runs[0].bold = True
    for (title, _, bullets), when in zip(EXPERIENCE, ("Recently", "Previously")):
        doc.add_paragraph(f"{title}    {when}")
        for bullet in bullets:
            doc.add_paragraph(bullet, style="List Bullet")
    doc.add_paragraph("EDUCATION").runs[0].bold = True
    doc.add_paragraph("B.S. Computer Science, University of Texas at Austin")
    doc.add_paragraph("TECHNICAL SKILLS").runs[0].bold = True
    doc.add_paragraph(SKILLS)
    doc.save(str(path))
    return ["parseable_dates"]


def build_sparse(path: Path) -> List[str]:
    """Almost no parseable text, the signature of a design-heavy resume."""
    doc = Document()
    doc.add_paragraph("JORDAN REYES")
    doc.add_paragraph("Data Engineer")
    doc.save(str(path))
    # A two-line document is missing everything, not just length. Declaring
    # only the headline defect would score the detector's correct findings as
    # false positives.
    return [
        "document_length", "contact_details", "parseable_dates", "section_headings",
    ]


#: Fixture name to (builder, defects it is built to contain).
BUILDERS: Dict[str, Callable[[Path], List[str]]] = {
    "clean": build_clean,
    "header_contact": build_header_contact,
    "table_layout": build_table_layout,
    "unlabeled_sections": build_unlabeled_sections,
    "no_dates": build_no_dates,
    "sparse": build_sparse,
}


def build_all(directory: str | Path) -> Dict[str, List[str]]:
    """Generate every fixture, returning each one's expected defects."""
    target = Path(directory)
    target.mkdir(parents=True, exist_ok=True)
    expected: Dict[str, List[str]] = {}
    for name, builder in BUILDERS.items():
        expected[name] = builder(target / f"{name}.docx")
    (target / "job_posting.txt").write_text(JOB_POSTING, encoding="utf-8")
    return expected
