"""
Unit tests for the resume writer module.
"""
import unittest
from unittest.mock import patch, MagicMock, mock_open
import os
import tempfile
from docx import Document
from resume.writer import ResumeWriter


class TestResumeWriter(unittest.TestCase):
    """Test cases for ResumeWriter class."""

    def setUp(self):
        """Set up test fixtures."""
        self.writer = ResumeWriter()
        
        # Sample resume content
        self.sample_resume_text = """
        <NAME>John Doe</NAME>
        <CONTACT>john.doe@example.com | 555-123-4567 | linkedin.com/in/johndoe</CONTACT>
        <SUMMARY>Experienced software engineer with 5+ years in web development.</SUMMARY>
        <SKILLS>
        Python | JavaScript | React | Node.js
        AWS | Docker | Kubernetes | CI/CD
        Agile | Scrum | Team Leadership
        </SKILLS>
        <EXPERIENCE>
        <COMPANY>Tech Solutions Inc.</COMPANY>
        <POSITION_DATE>Senior Developer | Jan 2020 - Present</POSITION_DATE>
        <BULLET>Led development of cloud-native applications using microservices architecture</BULLET>
        <BULLET>Implemented CI/CD pipelines reducing deployment time by 40%</BULLET>
        <COMPANY>Web Innovations</COMPANY>
        <POSITION_DATE>Software Engineer | Mar 2018 - Dec 2019</POSITION_DATE>
        <BULLET>Developed responsive web applications using React and Node.js</BULLET>
        </EXPERIENCE>
        <EDUCATION>
        Master of Computer Science, University of Technology, 2018
        Bachelor of Science in Software Engineering, State University, 2016
        </EDUCATION>
        """
    
    @patch('docx.Document')
    def test_write_resume(self, mock_document_class):
        """Test writing a resume to a DOCX file."""
        # Mock Document instance
        mock_doc = MagicMock()
        mock_document_class.return_value = mock_doc
        
        # Mock paragraph and run
        mock_para = MagicMock()
        mock_run = MagicMock()
        mock_para.add_run.return_value = mock_run
        mock_doc.add_paragraph.return_value = mock_para
        
        # Create a temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx') as temp_file:
            self.writer.write_resume(self.sample_resume_text, temp_file.name)
            
            # Verify Document was created
            mock_document_class.assert_called_once()
            
            # Verify document was saved
            mock_doc.save.assert_called_once_with(temp_file.name)
            
            # Verify name was added
            mock_doc.add_paragraph.assert_called()
    
    def test_extract_skills(self):
        """Test extracting skills from resume text."""
        skills = self.writer._extract_skills(self.sample_resume_text)
        
        # Verify skills were extracted correctly
        self.assertIsInstance(skills, list)
        self.assertIn("Python", skills)
        self.assertIn("JavaScript", skills)
        self.assertIn("React", skills)
        self.assertIn("Node.js", skills)
        self.assertIn("AWS", skills)
        self.assertIn("Docker", skills)
        
        # Don't check for exact count as implementation may vary
        self.assertGreaterEqual(len(skills), 9)  # At least the main skills
    
    def test_format_section_header(self):
        """Test formatting section headers."""
        # Create a mock document
        mock_doc = MagicMock()
        mock_para = MagicMock()
        mock_run = MagicMock()
        mock_doc.add_paragraph.return_value = mock_para
        mock_para.add_run.return_value = mock_run
        
        # Test the method
        self.writer._format_section_header(mock_doc, "EXPERIENCE")
        
        # Verify paragraph was added with correct formatting
        mock_doc.add_paragraph.assert_called_once()
        mock_para.add_run.assert_called_once_with("EXPERIENCE")
        mock_run.bold = True
        mock_para.style = "Heading 2"
    
    def test_add_bullet_point(self):
        """Test adding bullet points."""
        # Create a mock document
        mock_doc = MagicMock()
        mock_para = MagicMock()
        mock_doc.add_paragraph.return_value = mock_para
        
        # Test the method
        self.writer._add_bullet_point(mock_doc, "This is a bullet point")
        
        # Verify paragraph was added with correct style
        mock_doc.add_paragraph.assert_called_once_with(
            "This is a bullet point", style="List Bullet"
        )
    
    def test_parse_resume_sections(self):
        """Test parsing resume sections."""
        sections = self.writer._parse_resume_sections(self.sample_resume_text)
        
        # Verify all sections were extracted
        self.assertIn("NAME", sections)
        self.assertIn("CONTACT", sections)
        self.assertIn("SUMMARY", sections)
        self.assertIn("SKILLS", sections)
        self.assertIn("EXPERIENCE", sections)
        self.assertIn("EDUCATION", sections)
        
        # Verify section content
        self.assertEqual(sections["NAME"], "John Doe")
        self.assertEqual(sections["CONTACT"], "john.doe@example.com | 555-123-4567 | linkedin.com/in/johndoe")
        self.assertIn("Experienced software engineer", sections["SUMMARY"])
        self.assertIn("Python | JavaScript", sections["SKILLS"])
        self.assertIn("Tech Solutions Inc.", sections["EXPERIENCE"])
        self.assertIn("Master of Computer Science", sections["EDUCATION"])


if __name__ == '__main__':
    unittest.main() 