from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from docx import Document
from langchain_openai import ChatOpenAI  # Updated import
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate

import re
from docx.enum.text import WD_TAB_ALIGNMENT
import shutil
from openai import OpenAI
import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Twips
from docx.enum.style import WD_STYLE_TYPE
import os
import re
import time
import logging
from recommendations import generate_recommendations


# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def parse_date(date_str):
    """
    Parse a date string in MM/YYYY format and return a tuple (year, month) for sorting.
    """
    try:
        if '/' in date_str:
            month, year = date_str.strip().split('/')
            return (int(year), int(month))
        else:
            # Handle case where only year is provided
            return (int(date_str.strip()), 1)
    except (ValueError, IndexError):
        # Return a default value for unparseable dates
        logging.warning(f"Could not parse date: {date_str}")
        return (0, 0)

def add_section_heading(doc, text):
    """Add a section heading with consistent formatting and reduced spacing"""
    heading = doc.add_paragraph(style='Compact')
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_run = heading.add_run(text)
    heading_run.bold = True
    heading_run.font.size = Pt(12)
    heading_run.font.name = "Calibri"
    heading_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Add a horizontal line below the heading
    p = heading._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)
    
    # Reduce space after heading
    heading.paragraph_format.space_after = Pt(3)
    
    return heading

def add_bullet_point(doc, text, indent_level=0):
    """Add a bullet point with consistent formatting and reduced spacing"""
    bullet = doc.add_paragraph(style='Compact')  # Use compact style
    bullet.paragraph_format.left_indent = Inches(0.25 + (0.25 * indent_level))
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    
    # Check if the bullet point has a skill prefix (skill: text)
    if ':' in text:
        skill, achievement = text.split(':', 1)
        skill = skill.strip()
        achievement = achievement.strip()
        
        # Add bullet character
        bullet_char = bullet.add_run("• ")
        bullet_char.font.size = Pt(11)
        bullet_char.font.name = "Calibri"
        
        # Add skill in bold (not capitalized)
        skill_run = bullet.add_run(f"{skill.lower()}: ")
        skill_run.bold = True
        skill_run.font.size = Pt(11)
        skill_run.font.name = "Calibri"
        
        # Add achievement text
        achievement_run = bullet.add_run(achievement)
        achievement_run.font.size = Pt(11)
        achievement_run.font.name = "Calibri"
    else:
        # Add bullet character
        bullet_char = bullet.add_run("• ")
        bullet_char.font.size = Pt(11)
        bullet_char.font.name = "Calibri"
        
        # Just add the bullet text as is
        bullet_run = bullet.add_run(text)
        bullet_run.font.size = Pt(11)
        bullet_run.font.name = "Calibri"
    
    # Reduce space after bullet point
    bullet.paragraph_format.space_after = Pt(2)
    
    return bullet


def append_ok_to_sentences_with_full_formatting(input_path, output_path):
    """
    Loads a DOCX document from `input_path`, iterates over all paragraphs and their runs,
    and appends ' ok' to the end of each sentence while preserving the original formatting.
    The modified document is saved to `output_path`.
    """
    # Load the document
    doc = Document(input_path)

    # Iterate through all paragraphs
    for para in doc.paragraphs:
        if para.text.strip():  # Process only non-empty paragraphs
            new_runs = []  # Will store the new runs with modifications

            for run in para.runs:
                # Split sentences in the current run (preserving punctuation)
                sentences = run.text.split('. ')
                # Append " ok" to each sentence
                modified_sentences = [sentence if sentence else '' for sentence in sentences]
                modified_text = '. '.join(modified_sentences)

                # Create a new run with the same formatting as the original run
                new_run = para.add_run(modified_text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.strike = run.font.strike
                new_run.font.double_strike = run.font.double_strike
                new_run.font.all_caps = run.font.all_caps
                new_run.font.small_caps = run.font.small_caps
                new_run.font.shadow = run.font.shadow
                new_run.font.outline = run.font.outline
                new_run.font.emboss = run.font.emboss
                new_run.font.imprint = run.font.imprint
                new_run.font.hidden = run.font.hidden
                new_run.font.highlight_color = run.font.highlight_color
                if run.font.color:
                    new_run.font.color.rgb = run.font.color.rgb
                    new_run.font.color.theme_color = run.font.color.theme_color
                new_run.font.size = run.font.size
                new_run.font.name = run.font.name
                new_run.font.subscript = run.font.subscript
                new_run.font.superscript = run.font.superscript
                new_run.font.rtl = run.font.rtl
                new_run.font.complex_script = run.font.complex_script
                new_run.font.cs_bold = run.font.cs_bold
                new_run.font.cs_italic = run.font.cs_italic
                new_run.font.math = run.font.math
                new_run.font.no_proof = run.font.no_proof
                new_run.font.snap_to_grid = run.font.snap_to_grid
                new_run.font.spec_vanish = run.font.spec_vanish
                new_run.font.web_hidden = run.font.web_hidden

                new_runs.append(new_run)

            # Clear original runs and replace them with the new runs
            para.clear()
            for new_run in new_runs:
                para._element.append(new_run._element)

    # Save the modified document
    doc.save(output_path)

def process_resume(resume_path, job_description, output_path):
    """
    Processes the DOCX resume by appending 'ok' to the end of each sentence.
    The job_description parameter is kept for compatibility with the interface but is not used.
    """
    append_ok_to_sentences_with_full_formatting(resume_path, output_path)

def parse_markdown_line(line, paragraph):
    """
    Parses a markdown-formatted line and adds runs to the given paragraph.
    Supports **bold** and *italic* formatting.
    """
    # Split the line into tokens (captures **bold** and *italic*)
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
    for token in tokens:
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token.strip("**"))
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token.strip("*"))
            run.italic = True
        else:
            paragraph.add_run(token)

def reformat_resume_text_for_docx(revised_resume_text):
    """
    Reformats the revised resume text to be ATS-friendly and structured.
    - Detects markdown headings indicated by "##Heading##" (with or without spaces)
    - Applies "Heading 2" style to these headings.
    - Applies "List Bullet" style for bullet items starting with "-" or "*"
    - Applies "Heading 2" style for lines starting with a numeric prefix (e.g., "1. ")
    - Leaves all other lines as regular paragraphs.
    """
    new_doc = Document()
    for line in revised_resume_text.split("\n"):
        stripped_line = line.strip()
        if not stripped_line:
            continue
        # Handle markdown-style headings (e.g., "##AHMED DOGHRI##")
        if stripped_line.startswith("##") and stripped_line.endswith("##"):
            heading_text = stripped_line.strip("#").strip()
            p = new_doc.add_paragraph(heading_text, style="Heading 2")
        # Handle numeric headings like "1. Work Experience"
        elif re.match(r'^\d+\.\s', stripped_line):
            p = new_doc.add_paragraph(stripped_line, style="Heading 2")
        # Handle bullet list items starting with "-" or "*"
        elif stripped_line.startswith("-") or stripped_line.startswith("*"):
            p = new_doc.add_paragraph(stripped_line, style="List Bullet")
        else:
            # For other lines, check for date patterns to simulate right-indented dates.
            date_pattern = r'(\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}\s*-\s*(?:Present|\d{4}))'
            date_match = re.search(date_pattern, stripped_line, re.IGNORECASE)
            if date_match:
                main_text = stripped_line[:date_match.start()].strip()
                date_text = stripped_line[date_match.start():].strip()
                p = new_doc.add_paragraph()
                p.add_run(main_text + " ")
                # Add a right-aligned tab for the date text.
                tab_stop_position = 5000  # Adjust as necessary.
                p.paragraph_format.tab_stops.add_tab_stop(tab_stop_position)
                p.add_run("\t" + date_text)
            else:
                p = new_doc.add_paragraph(stripped_line)
    return new_doc



def format_resume_with_chatgpt(resume_text):
    """
    Uses GPT‑3.5 to reformat and reorder the resume text for maximum ATS compatibility.
    Instruct the model to output plain text with clear section headings, bold/italic markers,
    and proper ordering. Do not add new content.
    """
    format_prompt_template = (
        "You are an expert resume formatter. Reformat and reorder the following resume text "
        "to maximize its compatibility with Applicant Tracking Systems (ATS). Use clear section headings, "
        "apply bold or italics where appropriate, and ensure the resume is concise and well-structured. "
        "Output the final resume text with formatting markers (e.g., use '##' for section headings and wrap "
        "important keywords with ** for bold). Do not add any content that is not already present.\n\n"
        "Resume Text:\n{resume_text}\n\n"
        "Formatted Resume Text:"
    )
    format_prompt = PromptTemplate(
        template=format_prompt_template,
        input_variables=["resume_text"]
    )
    formatter = format_prompt | ChatOpenAI(model_name="gpt-3.5-turbo", temperature=0.1)
    input_data = {"resume_text": resume_text}
    formatted_result = formatter.invoke(input=input_data)
    return formatted_result.content



def rewrite_resume(
    resume_path: str, 
    job_description: str, 
    output_path: str,
    model: str = "gpt-4o",
    temperature: float = 0.1,
    api_key: str = None
):
    """
    Rewrites a resume to better match a job description using OpenAI's API.
    
    Args:
        resume_path: Path to the original resume DOCX file
        job_description: Text of the job posting
        output_path: Path where the rewritten resume will be saved
        model: OpenAI model to use
        temperature: Creativity level (0.0 to 1.0)
        api_key: OpenAI API key
    """
    try:
        # Extract text from the resume
        doc = Document(resume_path)
        resume_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        # Create OpenAI client
        client = OpenAI(api_key=api_key)
        
        # Improved prompt with specific steps and formatting instructions
        prompt = f"""
        You are an expert resume writer with years of experience helping job seekers land interviews.
        
        I'll provide you with a job description and a resume. Your task is to rewrite the resume to maximize the candidate's chances of getting an interview.
        
        Follow these specific steps:
        
        STEP 1: Identify the top 6 skills/requirements from the job description that the employer values most and name them to maximize keyword matching and recruiter engagement.
        
        STEP 2: Identify relevant positions and work experiences from the original resume that match those top skills.
        
        STEP 3: Rewrite each work experience to:
        - Emphasize achievements that demonstrate the top skills
        - Use keywords from the job description
        - Quantify achievements with metrics where possible
        - Optimize for ATS (Applicant Tracking Systems)
        - Make reasonable enhancements to maximize keyword matching (without fabricating major qualifications)
        - Write the work experience in reverse chronological order, with the most recent experience first
        - Format each bullet point as: "skill: Achievement with metric" (where skill is one of the top skills in lowercase)
        
        STEP 4: Structure the resume with these EXACT section markers (DO NOT include these markers in your output):
        - <NAME>: For the candidate's name
        - <CONTACT>: For contact information (phone, email, LinkedIn, etc.)
        - <SUMMARY>: For professional summary
        - <SKILLS>: For skills section (arrange the 6 skills in 3 lines with 2 skills per line)
        - <EXPERIENCE>: For work experience section
        - <EDUCATION>: For education section
        - <COMPANY>: For each company name
        - <POSITION_DATE>: For job title and dates (format as "Job Title | Dates")
        - <BULLET>: For each bullet point
        
        YOU MUST INCLUDE ALL OF THESE SECTION MARKERS IN YOUR RESPONSE. Do not skip any sections.
        
        STEP 5: Ensure the resume is comprehensive but concise:
        - Maximum 2 pages
        - Utilize the space effectively
        - Remove irrelevant information
        - Prioritize recent and relevant experiences
        - Prioritize skills that are most relevant to the job description
        - Prioritize work experiences over other sections like volunteerism and extracurricular activities
        - Prioritize longer work experiences over shorter ones, especially when there are multiple experiences around the same time period
        - Highlight promotions or achievements in the bullet points
        
        IMPORTANT: Use the section markers ONLY to indicate the structure. I will parse your response and remove these markers. The final resume should not contain any visible markers.
        
        JOB DESCRIPTION:
        {job_description}
        
        ORIGINAL RESUME:
        {resume_text}
        """
        
        # Check if using o1 or o3 models which have different API requirements
        if model.startswith('o1') or model.startswith('o3'):
            # For o1/o3 models, use a single user message without system message
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
        else:
            # For standard GPT models, use system+user message format
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You are an expert resume writer who helps job seekers optimize their resumes for specific job applications."},
                    {"role": "user", "content": prompt}
                ],
                temperature=temperature,
                max_tokens=4000
            )
        
        # Extract the rewritten resume text
        rewritten_resume = response.choices[0].message.content.strip()
        print("--------------------------------")
        print("REWRITTEN RESUME")
        print(rewritten_resume)
        print("--------------------------------")
        
        # Verify all required sections are present
        required_sections = [
            '<NAME>', '<CONTACT>', '<SUMMARY>', '<SKILLS>', 
            '<EXPERIENCE>', '<EDUCATION>', '<COMPANY>', 
            '<POSITION_DATE>', '<BULLET>'
        ]
        
        missing_sections = [section for section in required_sections if section not in rewritten_resume]
        
        # If any sections are missing, try to fix the resume
        if missing_sections:
            print(f"Missing sections: {missing_sections}")
            
            # Create a repair prompt to fix the missing sections
            repair_prompt = f"""
            The resume you generated is missing the following required section markers:
            {', '.join(missing_sections)}
            
            Please rewrite the resume to include ALL of these required section markers:
            - <NAME>: For the candidate's name
            - <CONTACT>: For contact information (phone, email, LinkedIn, etc.)
            - <SUMMARY>: For professional summary
            - <SKILLS>: For skills section (arrange the 6 skills in 3 lines with 2 skills per line)
            - <EXPERIENCE>: For work experience section
            - <EDUCATION>: For education section
            - <COMPANY>: For each company name
            - <POSITION_DATE>: For job title and dates (format as "Job Title | Dates")
            - <BULLET>: For each bullet point
            
            Here is the resume you generated:
            {rewritten_resume}
            
            for this job description:
            {job_description}
            
            Please rewrite it to include ALL the required section markers.
            """
            
            # Try to repair the resume
            if model.startswith('o1') or model.startswith('o3'):
                repair_response = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "user", "content": repair_prompt}
                    ],
                    temperature=1.0,
                )
            else:
                repair_response = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": "You are an expert resume writer who helps job seekers optimize their resumes for specific job applications."},
                        {"role": "user", "content": repair_prompt}
                    ],
                    temperature=temperature,
                    max_tokens=4000
                )
            
            # Use the repaired resume
            rewritten_resume = repair_response.choices[0].message.content.strip()
            print("--------------------------------")
            print("REPAIRED RESUME")
            print(rewritten_resume)
            print("--------------------------------")
            
            # Check if we still have missing sections
            still_missing = [section for section in required_sections if section not in rewritten_resume]
            if still_missing:
                print(f"Still missing sections after repair: {still_missing}")
                # Add placeholder sections as a last resort
                for section in still_missing:
                    if section == '<NAME>':
                        # Fix: Use a different approach to get the first line
                        first_line = resume_text.split('\n')[0] if '\n' in resume_text else resume_text
                        rewritten_resume = f"<NAME>: {first_line}\n" + rewritten_resume
                    elif section == '<CONTACT>':
                        rewritten_resume = rewritten_resume + f"\n<CONTACT>: Contact information from original resume"
                    elif section == '<SUMMARY>':
                        rewritten_resume = rewritten_resume + f"\n<SUMMARY>: Professional summary from original resume"
                    elif section == '<SKILLS>':
                        rewritten_resume = rewritten_resume + f"\n<SKILLS>: Skills from original resume"
                    elif section == '<EXPERIENCE>':
                        rewritten_resume = rewritten_resume + f"\n<EXPERIENCE>: Work experience from original resume"
                    elif section == '<EDUCATION>':
                        rewritten_resume = rewritten_resume + f"\n<EDUCATION>: Education from original resume"
                    elif section == '<COMPANY>':
                        rewritten_resume = rewritten_resume + f"\n<COMPANY>: Company from original resume"
                    elif section == '<POSITION_DATE>':
                        rewritten_resume = rewritten_resume + f"\n<POSITION_DATE>: Position and dates from original resume"
                    elif section == '<BULLET>':
                        rewritten_resume = rewritten_resume + f"\n<BULLET>: Achievement from original resume"
        
        
        
        # Create a new document with the rewritten content
        new_doc = Document()
        # Set document margins for better space utilization
        sections = new_doc.sections
        for section in sections:
            section.top_margin = Inches(0.3)
            section.bottom_margin = Inches(0.3)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Add a style for paragraphs with reduced spacing
        styles = new_doc.styles
        style = styles.add_style('Compact', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        # Process the resume text
        lines = rewritten_resume.strip().split('\n')
        
        # Extract name and contact info
        name = ""
        contact = ""
        for i, line in enumerate(lines):
            if line.strip() == "<NAME>":
                name = lines[i+1].strip()
            if line.strip() == "<CONTACT>":
                contact = lines[i+1].strip()
                break
        
        # Add name and contact info to document
        name_para = new_doc.add_paragraph(style='Compact')
        name_run = name_para.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(16)
        name_run.font.name = "Calibri"

        # Add contact info
        contact_para = new_doc.add_paragraph(style='Compact')
        contact_run = contact_para.add_run(contact)
        contact_run.font.size = Pt(11)
        contact_run.font.name = "Calibri"
        
        # Add space after header
        spacing_para = new_doc.add_paragraph(style='Compact')
        spacing_para.paragraph_format.space_after = Pt(2)
        
        # Extract summary
        summary = ""
        for i, line in enumerate(lines):
            if line.strip() == "<SUMMARY>":
                summary = lines[i+1].strip()
                break
        
        # Add executive summary section
        add_section_heading(new_doc, "EXECUTIVE SUMMARY")
        summary_para = new_doc.add_paragraph(style='Compact')
        summary_run = summary_para.add_run(summary)
        summary_run.font.size = Pt(11)
        summary_run.font.name = "Calibri"
        
        # Extract skills
        skills = []
        in_skills_section = False
        for line in lines:
            if line.strip() == "<SKILLS>":
                in_skills_section = True
                continue
            if in_skills_section and line.strip() and not line.startswith("<"):
                if "|" in line:
                    skills.extend([s.strip() for s in line.split("|")])
                else:
                    skills.append(line.strip())
            if in_skills_section and line.startswith("<"):
                in_skills_section = False
        
        # Add functional expertise section
        spacing_para = new_doc.add_paragraph(style='Compact')
        spacing_para.paragraph_format.space_after = Pt(2)
        add_section_heading(new_doc, "FUNCTIONAL EXPERTISE")
        
        # Calculate how many skills go in each column
        skills_per_column = (len(skills) + 1) // 2  # Ceiling division
        
        # Create skills in two columns using tab stops
        for i in range(skills_per_column):
            skills_para = new_doc.add_paragraph(style='Compact')
            
            # Set tab stop for second column
            skills_para.paragraph_format.tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.LEFT)
            
            # Add first column bullet and skill
            bullet_run = skills_para.add_run("• ")
            bullet_run.font.size = Pt(11)
            bullet_run.font.name = "Calibri"
            
            skill_run = skills_para.add_run(skills[i])
            skill_run.font.size = Pt(11)
            skill_run.font.name = "Calibri"
            
            # Add second column bullet and skill if available
            if i + skills_per_column < len(skills):
                tab_run = skills_para.add_run("\t")
                
                bullet_run2 = skills_para.add_run("• ")
                bullet_run2.font.size = Pt(11)
                bullet_run2.font.name = "Calibri"
                
                skill_run2 = skills_para.add_run(skills[i + skills_per_column])
                skill_run2.font.size = Pt(11)
                skill_run2.font.name = "Calibri"
        
        # Process experience sections
        current_company = None
        current_position = None
        current_date = None
        company_positions = {}  # Dictionary to track positions at each company

        # First pass: collect all positions for each company
        for i, line in enumerate(lines):
            if line.strip() == "<COMPANY>":
                current_company = lines[i+1].strip()
                if current_company not in company_positions:
                    company_positions[current_company] = []
            
            elif line.strip() == "<POSITION_DATE>" and current_company:
                position_date = lines[i+1].strip()
                if "|" in position_date:
                    current_position, current_date = [p.strip() for p in position_date.split("|", 1)]
                    
                    # Collect bullets for this position
                    bullets = []
                    j = i + 2
                    while j < len(lines):
                        if lines[j].strip() == "<BULLET>":
                            bullets.append(lines[j+1].strip())
                            j += 2
                        elif lines[j].strip() in ["<COMPANY>", "<POSITION_DATE>", "<EDUCATION>", "<VOLUNTEERISM>"]:
                            break
                        else:
                            j += 1
                    
                    # Add position with its date and bullets
                    company_positions[current_company].append((current_position, current_date, bullets))

        # Second pass: output companies with all their positions
        for company, positions in company_positions.items():
            # Add company and location
            company_para = new_doc.add_paragraph(style='Compact')
            
            if "|" in company:
                company_name, location = [c.strip() for c in company.split("|", 1)]
                company_run = company_para.add_run(f"{company_name} | {location}")
            else:
                company_run = company_para.add_run(company)
            
            company_run.bold = True
            company_run.font.size = Pt(12)
            company_run.font.name = "Calibri"
            
            # Set tab stop for date alignment
            tab_position = Inches(8.5 - 0.5)  # Page width minus right margin
            company_para.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT)
            
            # If multiple positions, add date range spanning all positions
            if len(positions) > 1:
                logging.info(f"Processing multiple positions for company: {company}")
                logging.info(f"Positions before sorting: {positions}")
                
                # Sort positions by date (assuming format MM/YYYY – MM/YYYY)
                positions.sort(key=lambda x: parse_date(x[1].split("–")[0].strip()), reverse=True)
                logging.info(f"Positions after sorting: {positions}")
                
                # Fix: Parse dates more carefully to handle edge cases
                all_start_dates = []
                all_end_dates = []
                all_start_tuples = []  # For accurate sorting
                
                logging.info("Extracting start and end dates from each position:")
                for position, date_range, _ in positions:
                    logging.info(f"  Position: {position}, Date range: {date_range}")
                    if "–" in date_range:
                        parts = date_range.split("–")
                        if len(parts) == 2:
                            start_date = parts[0].strip()
                            end_date = parts[1].strip()
                            all_start_dates.append(start_date)
                            all_end_dates.append(end_date)
                            all_start_tuples.append((start_date, parse_date(start_date)))
                            logging.info(f"    Extracted start date: {start_date}, end date: {end_date}")
                        else:
                            logging.warning(f"    Invalid date format: {date_range}")
                    else:
                        logging.warning(f"    No date separator found in: {date_range}")
                
                # Find earliest start date and latest end date
                logging.info(f"All start dates: {all_start_dates}")
                logging.info(f"All end dates: {all_end_dates}")
                
                if all_start_dates and all_end_dates:
                    # Sort dates using the parse_date function
                    all_start_tuples.sort(key=lambda x: x[1])  # Sort by the parsed tuple
                    earliest_date = all_start_tuples[0][0]  # Get the original date string
                    
                    # Sort end dates
                    end_date_tuples = [(date, parse_date(date)) for date in all_end_dates]
                    end_date_tuples.sort(key=lambda x: x[1], reverse=True)
                    latest_date = end_date_tuples[0][0]  # Get the original date string
                    
                    date_range = f"{earliest_date} – {latest_date}"
                    logging.info(f"Final date range: {date_range}")
                    
                    date_run = company_para.add_run("\t" + date_range)
                    date_run.bold = True  # Make company-level date bold
                    date_run.font.size = Pt(11)
                    date_run.font.name = "Calibri"
                else:
                    # Fallback to original method if parsing fails
                    logging.warning("Using fallback method for date extraction")
                    earliest_date = positions[-1][1].split("–")[0].strip()
                    latest_date = positions[0][1].split("–")[1].strip()
                    
                    logging.info(f"Fallback earliest date: {earliest_date}")
                    logging.info(f"Fallback latest date: {latest_date}")
                    
                    date_range = f"{earliest_date} – {latest_date}"
                    logging.info(f"Fallback date range: {date_range}")
                    
                    date_run = company_para.add_run("\t" + date_range)
                    date_run.bold = True  # Make company-level date bold
                    date_run.font.size = Pt(11)
                    date_run.font.name = "Calibri"
            else:
                # Just one position, add its date
                date_run = company_para.add_run("\t" + positions[0][1])
                date_run.bold = True  # Make company-level date bold
                date_run.font.size = Pt(11)
                date_run.font.name = "Calibri"
            
            # Add each position
            for position, date, bullets in positions:
                # Add position
                position_para = new_doc.add_paragraph(style='Compact')
                position_run = position_para.add_run(position)
                position_run.italic = True
                position_run.font.size = Pt(11)
                position_run.font.name = "Calibri"
                
                # Add date for individual position if multiple positions
                if len(positions) > 1:
                    position_para.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT)
                    date_run = position_para.add_run("\t" + date)
                    date_run.italic = True  # Keep italic
                    date_run.font.size = Pt(11)
                    date_run.font.name = "Calibri"
                
                # Add bullets for this position
                for bullet_text in bullets:
                    add_bullet_point(new_doc, bullet_text)
        
        # Add education section
        spacing_para = new_doc.add_paragraph(style='Compact')
        spacing_para.paragraph_format.space_after = Pt(2)
        add_section_heading(new_doc, "EDUCATION")
        
        # Extract education info
        education_lines = []
        in_education = False
        for line in lines:
            if line.strip() == "<EDUCATION>":
                in_education = True
                continue
            if in_education and line.strip() and not line.startswith("<"):
                education_lines.append(line.strip())
            if in_education and line.startswith("<") and line.strip() != "<EDUCATION>":
                in_education = False
        
        # Create education table
        if education_lines:
            for i in range(0, len(education_lines), 3):
                if i+2 < len(education_lines):
                    # Add institution and location with bullet
                    edu_para = new_doc.add_paragraph(style='Compact')
                    
                    # Add bullet
                    bullet_run = edu_para.add_run("• ")
                    bullet_run.font.size = Pt(11)
                    bullet_run.font.name = "Calibri"
                    
                    # Add institution and location
                    inst_run = edu_para.add_run(education_lines[i] + "  ")
                    inst_run.bold = True
                    inst_run.font.size = Pt(11)
                    inst_run.font.name = "Calibri"
                    
                    # Add date right-aligned
                    if "|" in education_lines[i+2]:
                        date = education_lines[i+2].split("|")[1].strip()
                        edu_para.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT)
                        date_run = edu_para.add_run("\t" + date)
                        date_run.bold = True  # Make education date bold
                        date_run.font.size = Pt(11)
                        date_run.font.name = "Calibri"
                    
                    # Add degree on next line with indentation
                    degree_para = new_doc.add_paragraph(style='Compact')
                    degree_para.paragraph_format.left_indent = Inches(0.25)
                    degree_run = degree_para.add_run(education_lines[i+1])
                    degree_run.font.size = Pt(11)
                    degree_run.font.name = "Calibri"
                    
                    # Add GPA on next line with indentation
                    if "GPA" in education_lines[i+2]:
                        gpa = education_lines[i+2].split("|")[0].strip() if "|" in education_lines[i+2] else education_lines[i+2]
                        gpa_para = new_doc.add_paragraph(style='Compact')
                        gpa_para.paragraph_format.left_indent = Inches(0.25)
                        gpa_run = gpa_para.add_run(gpa)
                        gpa_run.font.size = Pt(11)
                        gpa_run.font.name = "Calibri"
        
        # Add volunteerism section
        volunteerism_lines = []
        in_volunteerism = False
        for line in lines:
            if line.strip() == "<VOLUNTEERISM>":
                in_volunteerism = True
                continue
            if in_volunteerism and line.strip() and not line.startswith("<"):
                volunteerism_lines.append(line.strip())
            if in_volunteerism and line.startswith("<") and line.strip() != "<VOLUNTEERISM>":
                in_volunteerism = False
        
        if volunteerism_lines:
            spacing_para = new_doc.add_paragraph(style='Compact')
            spacing_para.paragraph_format.space_after = Pt(2)
            add_section_heading(new_doc, "VOLUNTEERISM")
            
            for i in range(0, len(volunteerism_lines), 2):
                if i+1 < len(volunteerism_lines):
                    # Add organization and date with bullet
                    vol_para = new_doc.add_paragraph(style='Compact')
                    
                    # Add bullet
                    bullet_run = vol_para.add_run("• ")
                    bullet_run.font.size = Pt(11)
                    bullet_run.font.name = "Calibri"
                    
                    # Add organization and date
                    if "|" in volunteerism_lines[i]:
                        org, date = [v.strip() for v in volunteerism_lines[i].split("|", 1)]
                        org_run = vol_para.add_run(org + "  ")
                        org_run.bold = True
                        org_run.font.size = Pt(11)
                        org_run.font.name = "Calibri"
                        
                        vol_para.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT)
                        date_run = vol_para.add_run("\t" + date)
                        date_run.bold = True  # Make volunteerism date bold
                        date_run.font.size = Pt(11)
                        date_run.font.name = "Calibri"
                    else:
                        org_run = vol_para.add_run(volunteerism_lines[i])
                        org_run.bold = True
                        org_run.font.size = Pt(11)
                        org_run.font.name = "Calibri"
                    
                    # Add description on next line with indentation
                    desc_para = new_doc.add_paragraph(style='Compact')
                    desc_para.paragraph_format.left_indent = Inches(0.25)
                    desc_run = desc_para.add_run(volunteerism_lines[i+1])
                    desc_run.font.size = Pt(11)
                    desc_run.font.name = "Calibri"
        
        # Add other relevant information section
        other_info_lines = []
        in_other_info = False
        for line in lines:
            if line.strip() == "<OTHER RELEVANT INFORMATION>":
                in_other_info = True
                continue
            if in_other_info and line.strip() and not line.startswith("<"):
                other_info_lines.append(line.strip())
            if in_other_info and line.startswith("<") and line.strip() != "<OTHER RELEVANT INFORMATION>":
                in_other_info = False

        if other_info_lines:
            spacing_para = new_doc.add_paragraph(style='Compact')
            spacing_para.paragraph_format.space_after = Pt(2)
            add_section_heading(new_doc, "OTHER RELEVANT INFORMATION")
            
            # Process each line as a separate bullet point
            for line in other_info_lines:
                if ":" in line:
                    # This line has a category and details
                    category, details = line.split(":", 1)
                    para = new_doc.add_paragraph(style='Compact')
                    para.paragraph_format.left_indent = Inches(0.25)
                    
                    # Add bullet
                    bullet_run = para.add_run("• ")
                    bullet_run.font.size = Pt(11)
                    bullet_run.font.name = "Calibri"
                    
                    # Add category in bold
                    cat_run = para.add_run(category.strip() + ": ")
                    cat_run.bold = True
                    cat_run.font.size = Pt(11)
                    cat_run.font.name = "Calibri"
                    
                    # Add details
                    details_run = para.add_run(details.strip())
                    details_run.font.size = Pt(11)
                    details_run.font.name = "Calibri"
                else:
                    # Just a regular line
                    para = new_doc.add_paragraph(style='Compact')
                    para.paragraph_format.left_indent = Inches(0.25)
                    
                    # Add bullet
                    bullet_run = para.add_run("• ")
                    bullet_run.font.size = Pt(11)
                    bullet_run.font.name = "Calibri"
                    
                    # Add text
                    text_run = para.add_run(line.strip())
                    text_run.font.size = Pt(11)
                    text_run.font.name = "Calibri"    

        # Save the document
        new_doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Resume rewriting failed: {e}. Falling back to original resume.")
        # If there's an error, copy the original resume to the output path
        if os.path.exists(resume_path):
            shutil.copy(resume_path, output_path)
        return False

def process_formatting(text):
    """
    Prepares text with formatting markers for processing.
    Returns a list of (text, is_bold, is_italic) tuples.
    """
    # Replace the markers with something we can parse more easily
    text = text.replace('**', '§BOLD§')
    text = text.replace('*', '§ITALIC§')
    
    # Split by formatting markers
    parts = []
    current_text = ""
    is_bold = False
    is_italic = False
    
    for char in text:
        if char == '§':
            if current_text:
                parts.append((current_text, is_bold, is_italic))
                current_text = ""
            
            # Check next few characters
            if text[text.index(char):].startswith('§BOLD§'):
                is_bold = not is_bold
                text = text[text.index(char) + 6:]  # Skip the marker
            elif text[text.index(char):].startswith('§ITALIC§'):
                is_italic = not is_italic
                text = text[text.index(char) + 8:]  # Skip the marker
        else:
            current_text += char
    
    if current_text:
        parts.append((current_text, is_bold, is_italic))
    
    return parts

def process_text_with_formatting(text, paragraph):
    """
    Adds text with proper formatting to a paragraph.
    """
    # Process bold text (between ** markers)
    bold_pattern = r'\*\*(.*?)\*\*'
    italic_pattern = r'\*(.*?)\*'
    
    # First, find all bold sections and process them
    bold_matches = re.finditer(bold_pattern, text)
    last_end = 0
    for match in bold_matches:
        # Add text before the bold section
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        
        # Add the bold text
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True
        
        last_end = match.end()
    
    # Add any remaining text after the last bold section
    if last_end < len(text):
        remaining_text = text[last_end:]
        
        # Process italic text in the remaining text
        italic_matches = re.finditer(italic_pattern, remaining_text)
        last_italic_end = 0
        
        for match in italic_matches:
            # Add text before the italic section
            if match.start() > last_italic_end:
                paragraph.add_run(remaining_text[last_italic_end:match.start()])
            
            # Add the italic text
            italic_run = paragraph.add_run(match.group(1))
            italic_run.italic = True
            
            last_italic_end = match.end()
        
        # Add any remaining text after the last italic section
        if last_italic_end < len(remaining_text):
            paragraph.add_run(remaining_text[last_italic_end:])